"""Handles exporting recipes to Word format."""

import io
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument  # actual type
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from miam.domain.entities import RecipeEntity
from miam.domain.ports_secondary import WordExporterPort


class WordExporter(WordExporterPort):
    """Secondary adapter that implements WordExporterPort."""

    def __init__(self, title: str = "My Recipe Book"):
        self.title = title  # Store title, not document

    def _create_fresh_document(self) -> DocxDocument:
        """Create a fresh document with title and styles."""
        doc = Document()
        self._setup_styles_for_doc(doc)
        self._add_title_to_doc(doc, self.title)
        return doc

    def _setup_styles_for_doc(self, doc: DocxDocument) -> None:
        """Configure document text styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_title_to_doc(self, doc: DocxDocument, title: str) -> None:
        """Add a centered heading at the document start."""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, recipes: list[RecipeEntity], output_path: str) -> None:
        """Export a list of RecipeEntity objects to a Word document."""
        self.document = self._create_fresh_document()
        for recipe in recipes:
            self._add_recipe(recipe)
            self.document.add_page_break()  # type: ignore[no-untyped-call] # ty: ignore[unused-ignore-comment]
        self.document.save(output_path)

    def to_bytes(self, recipes: list[RecipeEntity]) -> bytes:
        """Export a list of RecipeEntity objects to a Word document as bytes."""
        self.document = self._create_fresh_document()
        for recipe in recipes:
            self._add_recipe(recipe)
            self.document.add_page_break()  # type: ignore[no-untyped-call] # ty: ignore[unused-ignore-comment]

        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_recipe(self, recipe: RecipeEntity) -> None:
        # Title
        self.document.add_heading(recipe.title, level=1)

        # Meta info line
        meta_parts = [
            f"Category: {recipe.category}",
            f"Season: {recipe.season}" if recipe.season else None,
            "Vegetarian" if recipe.is_veggie else None,
            f"Difficulty: {recipe.difficulty}/3" if recipe.difficulty else None,
            f"Serves: {recipe.number_of_people}" if recipe.number_of_people else None,
            f"Rating: {recipe.rate}/5" if recipe.rate else None,
            "Tested" if recipe.tested else "Not tested",
        ]
        meta_text = " | ".join(p for p in meta_parts if p)

        meta_p = self.document.add_paragraph(meta_text)
        meta_p.runs[0].italic = True

        if recipe.tags:
            self.document.add_paragraph(f"Tags: {', '.join(recipe.tags)}")

        # Times table
        self._add_times_table(recipe)

        # Description
        self.document.add_heading("Description", level=2)
        self.document.add_paragraph(recipe.description)

        # Ingredients
        self.document.add_heading("Ingredients", level=2)
        for ing in recipe.ingredients:
            qty = f"{ing.quantity:g} " if ing.quantity else ""
            unit = f"{ing.unit} " if ing.unit else ""
            line = f"{qty}{unit}{ing.name}"
            self.document.add_paragraph(line, style="List Bullet")

        # Preparation steps
        if recipe.preparation:
            self.document.add_heading("Preparation", level=2)
            for i, step in enumerate(recipe.preparation, 1):
                self.document.add_paragraph(f"{i}. {step}")

    def _add_times_table(self, recipe: RecipeEntity) -> None:
        table = self.document.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Prep"
        hdr_cells[1].text = "Cook"
        hdr_cells[2].text = "Rest"

        row = table.add_row().cells
        row[0].text = self._fmt_minutes(recipe.prep_time_minutes)
        row[1].text = self._fmt_minutes(recipe.cook_time_minutes)
        row[2].text = self._fmt_minutes(recipe.rest_time_minutes)

    @staticmethod
    def _fmt_minutes(value: Any) -> str:
        return f"{value} min" if value is not None else "—"
