"""Tests for WordExporter producing real .docx documents."""

import io
import struct
import zlib
from pathlib import Path
from uuid import UUID, uuid4

from docx import Document

from miam.domain.entities import (
    ImageEntity,
    IngredientEntity,
    RecipeEntity,
    SourceEntity,
)
from miam.domain.ports_secondary import ImageStoragePort
from miam.domain.schemas import ImageResponse
from miam.infra.exporter_word import WordExporter


def _make_recipe(
    *,
    recipe_id: UUID | None = None,
    title: str = "Test Cake",
    description: str = "A delicious cake",
    category: str = "dessert",
    season: str | None = "summer",
    is_veggie: bool = True,
    difficulty: int | None = 2,
    number_of_people: int | None = 4,
    rate: int | None = 5,
    tested: bool = True,
    prep_time_minutes: int | None = 10,
    cook_time_minutes: int | None = 30,
    rest_time_minutes: int | None = 5,
    tags: list[str] | None = None,
    preparation: list[str] | None = None,
    ingredients: list[IngredientEntity] | None = None,
    images: list[ImageEntity] | None = None,
    sources: list[SourceEntity] | None = None,
) -> RecipeEntity:
    """Helper to build a RecipeEntity with sensible defaults."""
    return RecipeEntity(
        id=recipe_id or uuid4(),
        title=title,
        description=description,
        category=category,
        season=season,
        is_veggie=is_veggie,
        difficulty=difficulty,
        number_of_people=number_of_people,
        rate=rate,
        tested=tested,
        prep_time_minutes=prep_time_minutes,
        cook_time_minutes=cook_time_minutes,
        rest_time_minutes=rest_time_minutes,
        tags=tags if tags is not None else ["sweet"],
        preparation=preparation
        if preparation is not None
        else ["Mix ingredients", "Bake at 180C"],
        ingredients=ingredients
        if ingredients is not None
        else [
            IngredientEntity(name="Flour", quantity=200, unit="g"),
            IngredientEntity(name="Sugar", quantity=100, unit="g"),
        ],
        images=images if images is not None else [],
        sources=sources
        if sources is not None
        else [
            SourceEntity(type="manual", raw_content="Grandma's recipe"),
        ],
    )


def _read_docx_text(data: bytes) -> str:
    """Open docx bytes and return all paragraph + table cell text."""
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Save to file
# ---------------------------------------------------------------------------


class TestSave:
    def test_creates_file(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe()], str(out))
        assert out.is_file()

    def test_valid_docx(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe()], str(out))
        # python-docx can open it without error
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_contains_title(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe(title="Chocolate Cake")], str(out))
        text = _read_docx_text(out.read_bytes())
        assert "Chocolate Cake" in text

    def test_contains_ingredients(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe()], str(out))
        text = _read_docx_text(out.read_bytes())
        assert "Flour" in text
        assert "Sugar" in text

    def test_contains_preparation_steps(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe()], str(out))
        text = _read_docx_text(out.read_bytes())
        assert "Mix ingredients" in text
        assert "Bake at 180C" in text

    def test_contains_times_table(self, tmp_path: Path) -> None:
        exporter = WordExporter()
        out = tmp_path / "recipes.docx"
        exporter.save([_make_recipe()], str(out))
        text = _read_docx_text(out.read_bytes())
        assert "Prep" in text
        assert "Cook" in text
        assert "Rest" in text
        assert "10 min" in text
        assert "30 min" in text


# ---------------------------------------------------------------------------
# to_bytes
# ---------------------------------------------------------------------------


class TestToBytes:
    def test_returns_valid_docx_bytes(self) -> None:
        exporter = WordExporter()
        data = exporter.to_bytes([_make_recipe()])
        # Valid docx starts with PK (zip)
        assert data[:2] == b"PK"
        # python-docx can open it
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# With images
# ---------------------------------------------------------------------------


def _make_valid_png() -> bytes:
    """Generate a valid 1x1 white PNG."""
    raw_data = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw_data)

    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    png = b"\x89PNG\r\n\x1a\n"
    png += chunk(b"IHDR", ihdr_data)
    png += chunk(b"IDAT", compressed)
    png += chunk(b"IEND", b"")
    return png


class _StubImageStorage(ImageStoragePort):
    """Minimal stub implementing ImageStoragePort for word exporter tests."""

    def __init__(self, images: dict[UUID, ImageResponse] | None = None) -> None:
        self.images = images or {}

    def add_recipe_image(
        self, recipe_id: UUID, image: bytes, filename: str, image_id: UUID
    ) -> UUID:
        return image_id

    def get_recipe_image(self, image_id: UUID) -> ImageResponse | None:
        return self.images.get(image_id)

    def delete_image(self, image_id: UUID) -> bool:
        return image_id in self.images


class TestWithImages:
    def test_with_image_storage(self) -> None:
        image_id = uuid4()
        png_bytes = _make_valid_png()
        storage = _StubImageStorage(
            {image_id: ImageResponse(media_type="image/png", content=png_bytes)}
        )
        recipe = _make_recipe(images=[ImageEntity(id=image_id, caption="Photo")])
        exporter = WordExporter(image_storage=storage)
        data = exporter.to_bytes([recipe])

        text = _read_docx_text(data)
        assert "Photo" in text

    def test_without_image_storage(self) -> None:
        recipe = _make_recipe(images=[ImageEntity(id=uuid4(), caption="Photo")])
        exporter = WordExporter()
        data = exporter.to_bytes([recipe])
        # Should not crash
        assert data[:2] == b"PK"

    def test_missing_image_in_storage(self) -> None:
        storage = _StubImageStorage({})
        recipe = _make_recipe(images=[ImageEntity(id=uuid4(), caption="Missing")])
        exporter = WordExporter(image_storage=storage)
        data = exporter.to_bytes([recipe])
        # Should skip the missing image gracefully
        assert data[:2] == b"PK"


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


class TestEdgeCases:
    def test_empty_recipe_list(self) -> None:
        exporter = WordExporter()
        data = exporter.to_bytes([])
        doc = Document(io.BytesIO(data))
        # Only the document title heading
        assert len(doc.paragraphs) >= 1

    def test_recipe_without_preparation(self) -> None:
        recipe = _make_recipe(preparation=[])
        exporter = WordExporter()
        data = exporter.to_bytes([recipe])
        text = _read_docx_text(data)
        # Should not contain "Preparation" heading
        assert "Preparation" not in text
